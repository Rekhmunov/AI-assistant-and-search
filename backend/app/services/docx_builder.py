"""Сборка .docx из структуры DocumentStructure."""

from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.services.doc_gen_schema import DocumentStructure

LEGAL_DISCLAIMER = (
    "Документ подготовлен на основе общедоступных материалов и сформированных моделью данных. "
    "Он не является результатом юридической консультации и не заменяет помощь квалифицированного специалиста."
)

GLOSIX_FOOTER = "Подготовлено в Glosix"


def build_docx_bytes(
    structure: DocumentStructure,
    *,
    show_glosix_footer: bool,
) -> bytes:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    doc.add_heading(structure.title, level=0)

    for section in structure.sections:
        if section.heading:
            doc.add_heading(section.heading, level=1)
        for para in section.paragraphs:
            doc.add_paragraph(para)

    for table_def in structure.tables:
        if table_def.caption:
            cap = doc.add_paragraph(table_def.caption)
            cap.runs[0].bold = True
        col_count = max(len(table_def.headers), max((len(r) for r in table_def.rows), default=0))
        if col_count < 1:
            continue
        table = doc.add_table(rows=1, cols=col_count)
        table.style = "Table Grid"
        hdr_cells = table.rows[0].cells
        for i in range(col_count):
            hdr_cells[i].text = table_def.headers[i] if i < len(table_def.headers) else ""
        for row_data in table_def.rows:
            row = table.add_row().cells
            for i in range(col_count):
                row[i].text = row_data[i] if i < len(row_data) else ""
        doc.add_paragraph()

    doc.add_paragraph()
    disclaimer = doc.add_paragraph(LEGAL_DISCLAIMER)
    disclaimer.runs[0].italic = True

    if show_glosix_footer:
        footer_para = doc.add_paragraph(GLOSIX_FOOTER)
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
